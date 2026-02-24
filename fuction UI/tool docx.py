"""
title: Xuất DOCX (Wizard + chuẩn hoá Markdown → Word)
author: Thanh + AI
version: 1.0.0
required_open_webui_version: 0.6.0
requirements: python-docx,requests
"""

import asyncio
import base64
import datetime
import io
import logging
import re
from typing import Any, Dict, List, Optional

import requests
from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class Action:
    class Valves(BaseModel):
        # ---- UX ----
        show_status: bool = Field(default=True, description="Hiển thị status")
        show_confirmation: bool = Field(default=False, description="Hỏi xác nhận trước khi chạy")
        min_wizard_seconds: float = Field(default=1.5, description="Giữ wizard tối thiểu (giây)")

        # ---- OpenWebUI ----
        openwebui_base_url: str = Field(
            default="", description="Base URL Open WebUI (VD: http://localhost:3000). Để trống = tự suy."
        )
        openwebui_chat_completions_path: str = Field(
            default="/api/chat/completions", description="Path chat completions"
        )
        timeout_sec: int = Field(default=120, description="Timeout gọi API")

        # ---- Normalize scope ----
        normalize_mode: str = Field(
            default="conversation_window",
            description="assistant_only | conversation_window",
        )
        context_window_messages: int = Field(default=14, description="Số messages gần nhất")
        drop_file_like_content: bool = Field(
            default=True, description="Loại base64/data URL khỏi context"
        )

        # ---- Normalize prompt ----
        normalize_with_llm: bool = Field(default=True, description="Gọi LLM chuẩn hoá")
        normalize_instructions_vi: str = Field(
            default=(
                "Bạn là trợ lý chuẩn hoá nội dung.\n"
                "Nhiệm vụ: chuyển đổi hội thoại thành VĂN BẢN Markdown chuẩn để xuất Word.\n\n"
                "YÊU CẦU BẮT BUỘC:\n"
                "1) Trả về Markdown chuẩn, KHÔNG bọc trong ```.\n"
                "2) Dùng # cho tiêu đề chính, ## cho tiêu đề phụ, ### cho mục con.\n"
                "3) Dùng **bold**, *italic* cho nhấn mạnh.\n"
                "4) Dùng - hoặc * cho danh sách, 1. 2. 3. cho danh sách có thứ tự.\n"
                "5) Dùng bảng Markdown |col1|col2| nếu dữ liệu dạng bảng.\n"
                "6) Dùng > cho trích dẫn.\n"
                "7) Dùng ``` cho code blocks.\n"
                "8) Giữ nguyên nội dung có ý nghĩa, loại bỏ rác.\n"
                "9) Cấu trúc rõ ràng, chuyên nghiệp, phù hợp làm tài liệu.\n"
                "10) Nếu người dùng yêu cầu format cụ thể thì tuân thủ.\n"
            ),
            description="System prompt chuẩn hoá văn bản (VI)",
        )

        # ---- DOCX output ----
        base_filename: str = Field(default="van_ban.docx", description="Tên file")
        doc_title: str = Field(default="", description="Tiêu đề tài liệu (trống = tự lấy từ nội dung)")
        font_name: str = Field(default="Times New Roman", description="Font chữ")
        font_size_body: int = Field(default=12, description="Cỡ chữ body (pt)")
        font_size_h1: int = Field(default=18, description="Cỡ chữ H1 (pt)")
        font_size_h2: int = Field(default=15, description="Cỡ chữ H2 (pt)")
        font_size_h3: int = Field(default=13, description="Cỡ chữ H3 (pt)")
        page_margin_cm: float = Field(default=2.54, description="Margin trang (cm)")

    def __init__(self):
        self.valves = self.Valves()

    # =========================================================
    # UI MODAL - reuse pattern từ tool Excel
    # =========================================================

    async def _ui_modal_open(self, __event_call__=None):
        if not __event_call__:
            return
        js = """
(() => {
  const ID = "owui_export_docx_modal";
  if (document.getElementById(ID)) return;
  const overlay = document.createElement("div");
  overlay.id = ID;
  overlay.style.cssText = `
    position: fixed; inset: 0; z-index: 999999;
    background: rgba(0,0,0,0.45);
    display: flex; align-items: center; justify-content: center;
    font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial;
  `;
  overlay.innerHTML = `
    <div style="width: min(620px, 94vw); background: #fff; border-radius: 14px;
      box-shadow: 0 16px 60px rgba(0,0,0,.25); padding: 18px;">
      <div style="display:flex; gap:12px; align-items:flex-start;">
        <div id="owui_docx_spinner" style="width:34px;height:34px;border-radius:999px;
          border:3px solid #d1d5db;border-top-color:#7c3aed;
          animation:spin 0.9s linear infinite;margin-top:2px;"></div>
        <div style="flex:1">
          <div style="display:flex;justify-content:space-between;gap:12px;align-items:flex-start;">
            <div>
              <div style="font-size:16px;font-weight:700;color:#111827;" id="owui_docx_title">Đang xuất Word</div>
              <div style="font-size:13px;color:#4b5563;margin-top:4px;" id="owui_docx_subtitle">Khởi động...</div>
            </div>
            <div style="font-size:12px;color:#6b7280;margin-top:2px;"><span id="owui_docx_timer">0</span>s</div>
          </div>
          <div style="margin-top:12px;padding:10px;background:#f9fafb;border:1px solid #e5e7eb;border-radius:10px;">
            <div style="font-size:12px;color:#374151;font-weight:600;margin-bottom:6px;">Tiến trình</div>
            <ol style="margin:0;padding-left:18px;font-size:12.5px;color:#111827;line-height:1.65" id="owui_docx_steps">
              <li>Khởi tạo</li>
            </ol>
          </div>
          <div style="display:flex;justify-content:flex-end;gap:10px;margin-top:14px;">
            <button id="owui_docx_close" disabled
              style="opacity:.6;cursor:not-allowed;padding:8px 12px;border-radius:10px;border:1px solid #e5e7eb;background:#fff;">
              Đóng
            </button>
          </div>
          <div style="font-size:11px;color:#6b7280;margin-top:10px;">
            Thao tác có thể mất 10–60 giây tuỳ độ dài hội thoại.
          </div>
        </div>
      </div>
    </div>
  `;
  const style = document.createElement("style");
  style.textContent = `@keyframes spin { to { transform: rotate(360deg); } }`;
  overlay.appendChild(style);
  let sec = 0;
  const timerEl = overlay.querySelector("#owui_docx_timer");
  const t = setInterval(() => { sec++; if(timerEl) timerEl.textContent = String(sec); }, 1000);
  overlay.dataset.timer = String(t);
  document.body.appendChild(overlay);
})();
"""
        await __event_call__({"type": "execute", "data": {"code": js}})

    async def _ui_modal_update(self, __event_call__=None, title=None, subtitle=None,
                                steps=None, done_ok=None, error_text=None):
        if not __event_call__:
            return

        def _js(s):
            return s.replace("\\", "\\\\").replace("`", "\\`").replace("$", "\\$")

        steps_js = "null"
        if isinstance(steps, list):
            steps_js = "[" + ",".join([f"`{_js(x)}`" for x in steps]) + "]"
        t_js = f"`{_js(title)}`" if title else "null"
        st_js = f"`{_js(subtitle)}`" if subtitle else "null"
        e_js = f"`{_js(error_text)}`" if error_text else "null"
        d_js = "true" if done_ok is True else ("false" if done_ok is False else "null")

        js = f"""
(() => {{
  const o = document.getElementById("owui_export_docx_modal");
  if (!o) return;
  const t={t_js}, st={st_js}, steps={steps_js}, done={d_js}, err={e_js};
  if(t) o.querySelector("#owui_docx_title").textContent=t;
  if(st) o.querySelector("#owui_docx_subtitle").textContent=st;
  if(Array.isArray(steps)) o.querySelector("#owui_docx_steps").innerHTML=steps.map(s=>`<li>${{s}}</li>`).join("");
  if(done!==null) {{
    const tid=o.dataset.timer; if(tid) try{{clearInterval(Number(tid));}}catch(e){{}}
    const btn=o.querySelector("#owui_docx_close");
    btn.disabled=false;btn.style.opacity="1";btn.style.cursor="pointer";
    btn.onclick=()=>o.remove();
    const sp=o.querySelector("#owui_docx_spinner");
    if(sp){{sp.style.animation="none";
      sp.style.borderTopColor=done?"#16a34a":"#dc2626";
      sp.style.borderColor=done?"#86efac":"#fecaca";
      sp.style.background=done?"#dcfce7":"#fee2e2";}}
    if(done===true){{o.querySelector("#owui_docx_title").textContent="Xuất Word thành công";
      o.querySelector("#owui_docx_subtitle").textContent="File đang tải xuống.";}}
    else{{o.querySelector("#owui_docx_title").textContent="Xuất Word thất bại";
      o.querySelector("#owui_docx_subtitle").textContent=err||"Có lỗi xảy ra.";}}
  }}
}})();
"""
        await __event_call__({"type": "execute", "data": {"code": js}})

    # =========================================================
    # Context helpers (same as Excel tool)
    # =========================================================

    def _get_current_model_id(self, body, __model__=None):
        m = body.get("model")
        if isinstance(m, str) and m.strip():
            return m.strip()
        if isinstance(__model__, dict):
            for k in ("id", "model", "name"):
                v = __model__.get(k)
                if isinstance(v, str) and v.strip():
                    return v.strip()
        return None

    def _get_messages_window(self, body):
        msgs = body.get("messages", [])
        if not isinstance(msgs, list):
            return []
        n = int(self.valves.context_window_messages) if self.valves.context_window_messages else 0
        return msgs[-n:] if n > 0 else msgs

    def _sanitize_message_content(self, content):
        if content is None:
            return ""
        s = content if isinstance(content, str) else str(content)
        if self.valves.drop_file_like_content:
            if "data:" in s and "base64" in s:
                s = re.sub(r"data:[^\s]+;base64,[A-Za-z0-9+/=]+", "[DATA_URL_REMOVED]", s)
            s = re.sub(r"[A-Za-z0-9+/=]{1500,}", "[BASE64_REMOVED]", s)
        if len(s) > 6000:
            s = s[:6000] + "…[TRUNCATED]"
        return s

    def _build_normalize_user_prompt(self, body):
        mode = (self.valves.normalize_mode or "conversation_window").strip()
        if mode == "assistant_only":
            last_assistant = ""
            for m in reversed(body.get("messages", [])):
                if isinstance(m, dict) and m.get("role") == "assistant":
                    last_assistant = self._sanitize_message_content(m.get("content"))
                    if last_assistant.strip():
                        break
            return (
                "Hãy chuẩn hoá nội dung TRẢ LỜI CỦA BOT dưới đây thành văn bản Markdown chuyên nghiệp.\n\n"
                f"BOT:\n{last_assistant}"
            )
        window = self._get_messages_window(body)
        lines = []
        for m in window:
            if not isinstance(m, dict):
                continue
            role = m.get("role", "")
            c = self._sanitize_message_content(m.get("content"))
            if not c.strip():
                continue
            lines.append(f"[{role.upper()}]\n{c}")
        return (
            "Dưới đây là hội thoại. Hãy dựa vào YÊU CẦU CỦA NGƯỜI DÙNG để tạo văn bản Markdown "
            "chuyên nghiệp, có cấu trúc rõ ràng (heading, list, table, bold, italic...).\n\n"
            + "\n\n".join(lines)
        )

    # =========================================================
    # OpenWebUI API helpers
    # =========================================================

    def _resolve_base_url(self, __request__=None):
        if self.valves.openwebui_base_url.strip():
            return self.valves.openwebui_base_url.strip().rstrip("/")
        if __request__:
            try:
                base = str(__request__.base_url).rstrip("/")
                if base:
                    return base
            except Exception:
                pass
            try:
                origin = __request__.headers.get("origin")
                if origin:
                    return origin.rstrip("/")
                host = __request__.headers.get("host")
                proto = __request__.headers.get("x-forwarded-proto", "http")
                if host:
                    return f"{proto}://{host}".rstrip("/")
            except Exception:
                pass
        return None

    def _auth_headers(self, __request__=None):
        h = {"Content-Type": "application/json"}
        if not __request__:
            return h
        try:
            auth = __request__.headers.get("authorization")
            if auth:
                h["Authorization"] = auth
        except Exception:
            pass
        try:
            cookie = __request__.headers.get("cookie")
            if cookie and "Authorization" not in h:
                h["Cookie"] = cookie
        except Exception:
            pass
        return h

    async def _call_llm(self, body, __model__=None, __request__=None):
        base = self._resolve_base_url(__request__)
        if not base:
            raise ValueError("Không xác định được Open WebUI base_url.")
        url = base + self.valves.openwebui_chat_completions_path
        headers = self._auth_headers(__request__)
        model_id = self._get_current_model_id(body, __model__)
        if not model_id:
            raise ValueError("Không xác định được model.")
        user_prompt = self._build_normalize_user_prompt(body)
        messages = [
            {"role": "system", "content": self.valves.normalize_instructions_vi},
            {"role": "user", "content": user_prompt},
        ]
        payload = {"model": model_id, "messages": messages, "stream": False}

        def _do():
            return requests.post(url, headers=headers, json=payload,
                                 timeout=int(self.valves.timeout_sec))
        resp = await asyncio.to_thread(_do)
        if resp.status_code >= 400:
            raise ValueError(f"LLM call failed: HTTP {resp.status_code} - {resp.text[:500]}")
        data = resp.json()
        content = None
        try:
            content = data["choices"][0]["message"]["content"]
        except Exception:
            pass
        if not content:
            content = data.get("message", {}).get("content") or data.get("content") or ""
        if not content.strip():
            raise ValueError("LLM không trả về nội dung.")
        return content

    # =========================================================
    # Markdown → DOCX builder
    # =========================================================

    def _set_cell_shading(self, cell, color_hex):
        """Set background color for a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color_hex)
        tcPr.append(shading)

    def _add_formatted_run(self, paragraph, text, bold=False, italic=False, code=False):
        """Add a run with inline formatting."""
        run = paragraph.add_run(text)
        run.font.name = self.valves.font_name if not code else "Consolas"
        run.font.size = Pt(self.valves.font_size_body)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.size = Pt(self.valves.font_size_body - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        return run

    def _parse_inline(self, paragraph, text):
        """Parse inline Markdown: **bold**, *italic*, `code`, and plain text."""
        # Pattern for inline markdown
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'    # ***bold italic***
            r'|(\*\*(.+?)\*\*)'        # **bold**
            r'|(\*(.+?)\*)'            # *italic*
            r'|(`([^`]+)`)'            # `code`
            r'|([^*`]+)'              # plain text
        )
        for m in pattern.finditer(text):
            if m.group(2):  # bold italic
                self._add_formatted_run(paragraph, m.group(2), bold=True, italic=True)
            elif m.group(4):  # bold
                self._add_formatted_run(paragraph, m.group(4), bold=True)
            elif m.group(6):  # italic
                self._add_formatted_run(paragraph, m.group(6), italic=True)
            elif m.group(8):  # code
                self._add_formatted_run(paragraph, m.group(8), code=True)
            elif m.group(9):  # plain text
                self._add_formatted_run(paragraph, m.group(9))

    def _is_table_line(self, line):
        return line.count("|") >= 2 and len(line.strip()) >= 3

    def _is_separator_line(self, line):
        s = line.strip().strip("|").strip()
        return bool(s) and all(ch in "-: |" for ch in line.strip()) and "-" in line

    def _split_row(self, line):
        return [c.strip() for c in line.strip().strip("|").split("|")]

    def _build_docx_bytes(self, markdown_text: str) -> bytes:
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(self.valves.page_margin_cm)
            section.bottom_margin = Cm(self.valves.page_margin_cm)
            section.left_margin = Cm(self.valves.page_margin_cm)
            section.right_margin = Cm(self.valves.page_margin_cm)

        # Default font
        style = doc.styles["Normal"]
        style.font.name = self.valves.font_name
        style.font.size = Pt(self.valves.font_size_body)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        lines = markdown_text.splitlines()
        i = 0
        n = len(lines)
        in_code_block = False
        code_lines = []

        while i < n:
            line = lines[i]
            stripped = line.strip()

            # Code block start/end
            if stripped.startswith("```"):
                if in_code_block:
                    # End code block - render collected lines
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    for ci, cl in enumerate(code_lines):
                        run = p.add_run(cl)
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                        if ci < len(code_lines) - 1:
                            p.add_run("\n").font.size = Pt(10)
                    # Add light gray background via shading
                    pPr = p._p.get_or_add_pPr()
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:val"), "clear")
                    shading.set(qn("w:fill"), "F5F5F5")
                    pPr.append(shading)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Empty line
            if not stripped:
                i += 1
                continue

            # Headings
            if stripped.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[4:])
                run.bold = True
                run.font.size = Pt(self.valves.font_size_h3)
                run.font.name = self.valves.font_name
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                p.paragraph_format.space_before = Pt(12)
                i += 1
                continue

            if stripped.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[3:])
                run.bold = True
                run.font.size = Pt(self.valves.font_size_h2)
                run.font.name = self.valves.font_name
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                p.paragraph_format.space_before = Pt(14)
                i += 1
                continue

            if stripped.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:])
                run.bold = True
                run.font.size = Pt(self.valves.font_size_h1)
                run.font.name = self.valves.font_name
                run.font.color.rgb = RGBColor(0x0D, 0x37, 0x6B)
                p.paragraph_format.space_before = Pt(18)
                i += 1
                continue

            # Blockquote
            if stripped.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                self._parse_inline(p, stripped[2:])
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                i += 1
                continue

            # Horizontal rule
            if stripped in ("---", "***", "___"):
                p = doc.add_paragraph()
                p.add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                i += 1
                continue

            # Table (Markdown)
            if (self._is_table_line(stripped)
                    and i + 1 < n
                    and self._is_separator_line(lines[i + 1].strip())):
                # Collect table block
                headers = self._split_row(stripped)
                i += 2  # skip header + separator
                data_rows = []
                while i < n and self._is_table_line(lines[i].strip()):
                    data_rows.append(self._split_row(lines[i].strip()))
                    i += 1

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"

                # Header row
                for ci, h in enumerate(headers):
                    cell = table.rows[0].cells[ci] if ci < num_cols else None
                    if cell:
                        cell.text = h
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.bold = True
                                run.font.size = Pt(self.valves.font_size_body)
                                run.font.name = self.valves.font_name
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        self._set_cell_shading(cell, "1F4E79")

                # Data rows
                for ri, row in enumerate(data_rows):
                    for ci in range(num_cols):
                        cell = table.rows[ri + 1].cells[ci]
                        val = row[ci] if ci < len(row) else ""
                        cell.text = val
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(self.valves.font_size_body)
                                run.font.name = self.valves.font_name
                        # Alternate row shading
                        if ri % 2 == 0:
                            self._set_cell_shading(cell, "F2F7FB")
                continue

            # Unordered list
            if re.match(r'^[\-\*\+]\s', stripped):
                text = re.sub(r'^[\-\*\+]\s+', '', stripped)
                p = doc.add_paragraph(style="List Bullet")
                self._parse_inline(p, text)
                i += 1
                continue

            # Ordered list
            ol_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if ol_match:
                text = ol_match.group(2)
                p = doc.add_paragraph(style="List Number")
                self._parse_inline(p, text)
                i += 1
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            self._parse_inline(p, stripped)
            i += 1

        # Footer with timestamp
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Tạo lúc: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = self.valves.font_name

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    # =========================================================
    # Main action
    # =========================================================

    async def action(self, body, __user__=None, __event_emitter__=None,
                     __event_call__=None, __model__=None, __request__=None, __id__=None):
        wizard_start = asyncio.get_event_loop().time()
        steps = [
            "1) Chuẩn hoá nội dung thành Markdown",
            "2) Chuyển Markdown → file Word (.docx)",
            "3) Tải file xuống trình duyệt",
        ]

        try:
            if self.valves.show_confirmation and __event_call__:
                ok = await __event_call__({
                    "type": "confirmation",
                    "data": {"title": "Xuất Word", "message": "Bắt đầu xuất Word từ hội thoại?"},
                })
                if not ok:
                    if __event_emitter__:
                        await __event_emitter__({"type": "notification",
                            "data": {"type": "info", "content": "Đã huỷ xuất Word."}})
                    return body

            await self._ui_modal_open(__event_call__)
            await self._ui_modal_update(__event_call__, subtitle="Đang khởi tạo...", steps=steps)

            if __event_emitter__ and self.valves.show_status:
                await __event_emitter__({"type": "status",
                    "data": {"description": "Đang xuất Word...", "done": False}})

            # Step 1: Normalize
            await self._ui_modal_update(__event_call__,
                subtitle="Bước 1/3: Chuẩn hoá nội dung...", steps=steps)
            if self.valves.normalize_with_llm:
                markdown_text = await self._call_llm(body, __model__, __request__)
            else:
                # Fallback: lấy message cuối assistant
                markdown_text = ""
                for m in reversed(body.get("messages", [])):
                    if isinstance(m, dict) and m.get("role") == "assistant":
                        markdown_text = m.get("content", "")
                        if markdown_text.strip():
                            break

            # Step 2: Build DOCX
            await self._ui_modal_update(__event_call__,
                subtitle="Bước 2/3: Đang tạo file Word...", steps=steps)
            docx_bytes = self._build_docx_bytes(markdown_text)

            # Step 3: Download
            await self._ui_modal_update(__event_call__,
                subtitle="Bước 3/3: Đang tải file xuống...", steps=steps)

            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{ts}_{self.valves.base_filename}"
            b64 = base64.b64encode(docx_bytes).decode("utf-8")

            js = f"""
try {{
  const b = atob("{b64}");
  const bytes = new Uint8Array(b.length);
  for(let i=0;i<b.length;i++) bytes[i]=b.charCodeAt(i);
  const blob = new Blob([bytes],{{type:"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}});
  const url = URL.createObjectURL(blob);
  const a = document.createElement("a");
  a.href=url; a.download="{filename}";
  document.body.appendChild(a); a.click();
  URL.revokeObjectURL(url); a.remove();
}} catch(e) {{ console.error("Download DOCX failed",e); }}
"""
            if __event_call__:
                await __event_call__({"type": "execute", "data": {"code": js}})

            elapsed = asyncio.get_event_loop().time() - wizard_start
            if elapsed < self.valves.min_wizard_seconds:
                await asyncio.sleep(self.valves.min_wizard_seconds - elapsed)

            await self._ui_modal_update(__event_call__, done_ok=True, steps=steps)

            if __event_emitter__ and self.valves.show_status:
                await __event_emitter__({"type": "status",
                    "data": {"description": "Hoàn tất xuất Word.", "done": True}})
                await __event_emitter__({"type": "notification",
                    "data": {"type": "success", "content": "Đã xuất Word thành công."}})
            return body

        except Exception as e:
            logger.exception("Lỗi xuất DOCX: %s", str(e))
            elapsed = asyncio.get_event_loop().time() - wizard_start
            if elapsed < self.valves.min_wizard_seconds:
                await asyncio.sleep(self.valves.min_wizard_seconds - elapsed)
            await self._ui_modal_update(__event_call__, done_ok=False,
                error_text=str(e), steps=steps)
            if __event_emitter__:
                await __event_emitter__({"type": "status",
                    "data": {"description": f"Lỗi: {e}", "done": True}})
                await __event_emitter__({"type": "notification",
                    "data": {"type": "error", "content": f"Lỗi xuất Word: {e}"}})
            return None
