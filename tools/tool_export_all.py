"""
title: Xuất File (PDF / Excel / Word)
author: Thanh + AI
version: 1.0.0
required_open_webui_version: 0.6.0
requirements: fpdf2>=2.7,openpyxl>=3.1,python-docx>=1.1,requests>=2.28
"""

import asyncio
import base64
import datetime
import io
import logging
import os
import re
from typing import Any, Dict, List, Optional

import requests
from pydantic import BaseModel, Field

# --- PDF ---
from fpdf import FPDF

# --- Excel ---
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

# --- DOCX ---
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


# =====================================================================
# PDF helper class (from tool pdf.py)
# =====================================================================

class MarkdownPDF(FPDF):
    """Custom FPDF subclass for rendering Markdown to PDF."""

    def __init__(self, font_name="DejaVu", font_size=11, margin=15):
        super().__init__()
        self._font_name = font_name
        self._font_size = font_size
        self.set_auto_page_break(auto=True, margin=margin)
        self.set_margins(margin, margin, margin)

        # Register DejaVu Unicode font (cài qua apt: fonts-dejavu-core + fonts-dejavu-extra)
        # Sử dụng fname= với đường dẫn tuyệt đối theo tài liệu fpdf2:
        # https://py-pdf.github.io/fpdf2/Unicode.html#adding-and-using-fonts
        _fd = "/usr/share/fonts/truetype/dejavu"
        self.add_font("DejaVu", style="", fname=f"{_fd}/DejaVuSans.ttf")
        self.add_font("DejaVu", style="b", fname=f"{_fd}/DejaVuSans-Bold.ttf")
        self.add_font("DejaVu", style="i", fname=f"{_fd}/DejaVuSans-Oblique.ttf")
        self.add_font("DejaVu", style="bi", fname=f"{_fd}/DejaVuSans-BoldOblique.ttf")

    def header(self):
        pass

    def footer(self):
        self.set_y(-12)
        self.set_font(self._font_name, "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 8, f"Trang {self.page_no()}/{{nb}}", align="C")


class Action:
    class Valves(BaseModel):
        # ---- UX ----
        show_status: bool = Field(
            default=True, description="Hiển thị status/notification"
        )
        show_confirmation: bool = Field(
            default=False, description="Hỏi xác nhận trước khi chạy"
        )
        min_wizard_seconds: float = Field(
            default=1.5, description="Giữ wizard tối thiểu (giây)"
        )

        # ---- OpenWebUI internal call ----
        openwebui_base_url: str = Field(
            default="",
            description="Base URL Open WebUI (VD: http://localhost:3000). Để trống thì cố suy ra từ request.",
        )
        openwebui_chat_completions_path: str = Field(
            default="/api/chat/completions", description="Path chat completions"
        )
        timeout_sec: int = Field(
            default=120, description="Timeout gọi /api/chat/completions"
        )

        # ---- Normalize scope ----
        normalize_mode: str = Field(
            default="conversation_window",
            description="assistant_only | conversation_window",
        )
        context_window_messages: int = Field(
            default=14,
            description="Số messages gần nhất đưa vào normalize (khi conversation_window)",
        )
        drop_file_like_content: bool = Field(
            default=True,
            description="Loại các đoạn giống file/base64 khỏi context (giảm token)",
        )

        # ---- Normalize prompt ----
        normalize_with_llm: bool = Field(
            default=True, description="Gọi LLM để chuẩn hoá"
        )
        max_rows_hint: int = Field(
            default=250, description="Gợi ý giới hạn số dòng dữ liệu (Excel)"
        )

        # ---- Normalize instructions per format ----
        normalize_instructions_pdf: str = Field(
            default=(
                "Bạn là trợ lý chuẩn hoá nội dung.\n"
                "Nhiệm vụ: chuyển hội thoại thành VĂN BẢN Markdown chuẩn để xuất PDF.\n\n"
                "BƯỚC 1 — TỰ PHÁT HIỆN LOẠI NỘI DUNG:\n"
                "Đọc toàn bộ hội thoại, xác định loại nội dung chính:\n"
                "• Bảng/dữ liệu → ưu tiên bảng Markdown gọn, header rõ ràng\n"
                "• Hướng dẫn/tutorial → đánh số bước, heading phân cấp, code block nếu có\n"
                "• Phân tích/báo cáo → heading + bullet points + bảng tóm tắt\n"
                "• Hỏi đáp/Q&A → heading cho câu hỏi, nội dung trả lời bên dưới\n"
                "• So sánh/đánh giá → bảng so sánh + nhận xét\n"
                "• Sáng tạo/viết lách → giữ nguyên văn phong, format đẹp\n"
                "• Danh sách/liệt kê → bullet/numbered list rõ ràng\n"
                "• Tổng hợp/tóm tắt → heading + key points\n"
                "nói chung là tùy vào nội dung để đưa ra md phù hơp nhất\n\n"
                "BƯỚC 2 — CHUẨN HOÁ MARKDOWN:\n"
                "1) Trả về Markdown chuẩn, KHÔNG bọc trong ```.\n"
                "2) Dùng # cho tiêu đề chính, ## cho tiêu đề phụ.\n"
                "3) Dùng **bold**, *italic* cho nhấn mạnh.\n"
                "4) Dùng - hoặc * cho danh sách.\n"
                "5) Dùng bảng Markdown |col1|col2| nếu dữ liệu dạng bảng.\n"
                "6) Dùng > cho trích dẫn.\n"
                "7) Cấu trúc phù hợp với loại nội dung đã phát hiện ở Bước 1.\n"
                "8) Nếu người dùng yêu cầu cụ thể thì tuân thủ (ưu tiên cao nhất).\n"
            ),
            description="System prompt chuẩn hoá PDF (VI)",
        )
        normalize_instructions_excel: str = Field(
            default=(
                "Bạn là trợ lý chuẩn hoá dữ liệu.\n"
                "Nhiệm vụ: tạo ĐÚNG 1 bảng Markdown duy nhất từ nội dung hội thoại.\n\n"
                "BƯỚC 1 — TỰ PHÁT HIỆN LOẠI DỮ LIỆU:\n"
                "Đọc hội thoại, xác định dạng dữ liệu:\n"
                "• Bảng có sẵn → giữ cấu trúc, chuẩn hoá header\n"
                "• Danh sách/liệt kê → chuyển thành bảng có cột phù hợp\n"
                "• So sánh → tạo bảng so sánh (cột = tiêu chí, hàng = đối tượng)\n"
                "• Số liệu rải rác → gom thành bảng có header hợp lý\n"
                "• Q&A → bảng 2 cột: Câu hỏi | Trả lời\n"
                "• Timeline/lịch sử → bảng có cột thời gian\n"
                "nói chung là tùy vào nội dung để đưa ra md phù hơp nhất\n\n"
                "BƯỚC 2 — TẠO BẢNG MARKDOWN:\n"
                "1) Chỉ trả về bảng Markdown, KHÔNG thêm văn bản trước/sau, KHÔNG bọc ```.\n"
                "2) Dòng 1 là header, dòng 2 là separator |---|.\n"
                "3) Header ngắn gọn, rõ nghĩa, phù hợp loại dữ liệu.\n"
                "4) Không xuống dòng trong ô.\n"
                "5) Nếu thiếu dữ liệu, để trống ô.\n"
                "6) Không tạo nhiều bảng.\n"
                "7) Nếu người dùng yêu cầu lọc/cột/số dòng/format thì tuân thủ (ưu tiên cao nhất).\n"
            ),
            description="System prompt chuẩn hoá Excel (VI)",
        )
        normalize_instructions_docx: str = Field(
            default=(
                "Bạn là trợ lý chuẩn hoá nội dung.\n"
                "Nhiệm vụ: chuyển đổi hội thoại thành VĂN BẢN Markdown chuẩn để xuất Word.\n\n"
                "BƯỚC 1 — TỰ PHÁT HIỆN LOẠI NỘI DUNG:\n"
                "Đọc toàn bộ hội thoại, xác định loại nội dung chính:\n"
                "• Bảng/dữ liệu → bảng Markdown + heading mô tả\n"
                "• Hướng dẫn/tutorial → heading phân cấp, bước đánh số, code block\n"
                "• Phân tích/báo cáo → heading + bullet + bảng tóm tắt\n"
                "• Hỏi đáp/Q&A → heading cho câu hỏi, trả lời bên dưới\n"
                "• So sánh/đánh giá → bảng so sánh + nhận xét\n"
                "• Sáng tạo/viết lách → giữ văn phong, format heading + paragraph\n"
                "• Danh sách → bullet/numbered list\n"
                "• Tài liệu kỹ thuật → heading phân cấp, code block, bảng specs\n"
                "nói chung là tùy vào nội dung để đưa ra md phù hơp nhất\n\n"
                "BƯỚC 2 — CHUẨN HOÁ MARKDOWN:\n"
                "1) Trả về Markdown chuẩn, KHÔNG bọc trong ```.\n"
                "2) Dùng # cho tiêu đề chính, ## cho tiêu đề phụ, ### cho mục con.\n"
                "3) Dùng **bold**, *italic* cho nhấn mạnh.\n"
                "4) Dùng - hoặc * cho danh sách, 1. 2. 3. cho danh sách có thứ tự.\n"
                "5) Dùng bảng Markdown |col1|col2| nếu dữ liệu dạng bảng.\n"
                "6) Dùng > cho trích dẫn, ``` cho code blocks.\n"
                "7) Giữ nguyên nội dung có ý nghĩa, loại bỏ rác.\n"
                "8) Cấu trúc phù hợp với loại nội dung đã phát hiện ở Bước 1.\n"
                "9) Nếu người dùng yêu cầu format cụ thể thì tuân thủ (ưu tiên cao nhất).\n"
            ),
            description="System prompt chuẩn hoá DOCX (VI)",
        )

        # ---- PDF output ----
        pdf_base_filename: str = Field(default="van_ban.pdf", description="Tên file PDF")
        pdf_font_size_body: int = Field(default=11, description="Cỡ chữ body PDF (pt)")
        pdf_font_size_h1: int = Field(default=18, description="Cỡ H1 PDF")
        pdf_font_size_h2: int = Field(default=15, description="Cỡ H2 PDF")
        pdf_font_size_h3: int = Field(default=13, description="Cỡ H3 PDF")
        pdf_page_margin: int = Field(default=15, description="Margin trang PDF (mm)")

        # ---- Excel output ----
        xlsx_sheet_name: str = Field(default="Extracted", description="Tên sheet Excel")
        xlsx_base_filename: str = Field(
            default="trich_xuat.xlsx", description="Tên file Excel"
        )
        xlsx_title: str = Field(
            default="BẢNG TRÍCH XUẤT", description="Tiêu đề trên đầu file Excel"
        )
        xlsx_table_index: int = Field(
            default=0, description="Nếu output có nhiều bảng, chọn bảng N (0-based)"
        )
        xlsx_table_start_row: int = Field(
            default=5, description="Dòng bắt đầu bảng (header) trong Excel"
        )
        xlsx_max_col_width: int = Field(default=55, description="Độ rộng cột tối đa Excel")

        # ---- DOCX output ----
        docx_base_filename: str = Field(default="van_ban.docx", description="Tên file Word")
        docx_doc_title: str = Field(default="", description="Tiêu đề tài liệu Word")
        docx_font_name: str = Field(default="Times New Roman", description="Font chữ Word")
        docx_font_size_body: int = Field(default=12, description="Cỡ chữ body Word (pt)")
        docx_font_size_h1: int = Field(default=18, description="Cỡ H1 Word (pt)")
        docx_font_size_h2: int = Field(default=15, description="Cỡ H2 Word (pt)")
        docx_font_size_h3: int = Field(default=13, description="Cỡ H3 Word (pt)")
        docx_page_margin_cm: float = Field(default=2.54, description="Margin trang Word (cm)")

    def __init__(self):
        self.valves = self.Valves()

    # =========================================================
    # FORMAT SELECTION DIALOG
    # =========================================================

    async def _show_format_picker(self, __event_call__=None):
        """Show format selection. Returns 'pdf', 'excel', 'docx' or None."""
        if not __event_call__:
            return None

        result = await __event_call__(
            {
                "type": "input",
                "data": {
                    "title": "📦 Chọn định dạng xuất",
                    "message": "1 = 📊 Excel\n2 = 📄 PDF\n3 = 📝 Word\n\nNhập 1, 2, 3 hoặc tên định dạng:",
                    "placeholder": "1",
                },
            }
        )

        if not result or not isinstance(result, str):
            return None

        fmt = result.strip().lower()
        mapping = {
            "1": "excel", "2": "pdf", "3": "docx",
            "excel": "excel", "xlsx": "excel", "xls": "excel",
            "pdf": "pdf",
            "docx": "docx", "doc": "docx", "word": "docx",
        }
        return mapping.get(fmt)

    # =========================================================
    # UI MODAL (from Excel tool - used as-is)
    # =========================================================

    async def _ui_modal_open(self, __event_call__=None, fmt="excel"):
        if not __event_call__:
            return

        fmt_labels = {"pdf": "PDF", "excel": "Excel", "docx": "Word"}
        fmt_colors = {"pdf": "#dc2626", "excel": "#2563eb", "docx": "#7c3aed"}
        label = fmt_labels.get(fmt, "File")
        color = fmt_colors.get(fmt, "#2563eb")

        js = f"""
(() => {{
  const ID = "owui_export_modal";
  if (document.getElementById(ID)) return;

  const overlay = document.createElement("div");
  overlay.id = ID;
  overlay.style.cssText = `
    position: fixed; inset: 0; z-index: 999999;
    background: rgba(0,0,0,0.45);
    display: flex; align-items: center; justify-content: center;
    font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial;
  `;

  overlay.innerHTML = `
    <div style="
      width: min(620px, 94vw);
      background: #fff; border-radius: 14px;
      box-shadow: 0 16px 60px rgba(0,0,0,.25);
      padding: 18px 18px 14px 18px;
    ">
      <div style="display:flex; gap:12px; align-items:flex-start;">
        <div id="owui_modal_spinner" style="
          width: 34px; height: 34px; border-radius: 999px;
          border: 3px solid #d1d5db; border-top-color: {color};
          animation: spin 0.9s linear infinite; margin-top: 2px;
        "></div>

        <div style="flex:1">
          <div style="display:flex; justify-content:space-between; gap:12px; align-items:flex-start;">
            <div>
              <div style="font-size:16px; font-weight:700; color:#111827;" id="owui_modal_title">Đang xuất {label}</div>
              <div style="font-size:13px; color:#4b5563; margin-top:4px;" id="owui_modal_subtitle">Khởi động tiến trình...</div>
            </div>
            <div style="font-size:12px; color:#6b7280; margin-top:2px;">
              <span id="owui_modal_timer">0</span>s
            </div>
          </div>

          <div style="margin-top:12px; padding:10px 10px; background:#f9fafb; border:1px solid #e5e7eb; border-radius:10px;">
            <div style="font-size:12px; color:#374151; font-weight:600; margin-bottom:6px;">Tiến trình</div>
            <ol style="margin:0; padding-left:18px; font-size:12.5px; color:#111827; line-height:1.65" id="owui_modal_steps">
              <li>Khởi tạo tiến trình</li>
            </ol>
          </div>

          <div style="display:flex; justify-content:flex-end; gap:10px; margin-top:14px;">
            <button id="owui_modal_close" disabled
              style="opacity:.6; cursor:not-allowed; padding:8px 12px; border-radius:10px; border:1px solid #e5e7eb; background:#fff;">
              Đóng
            </button>
          </div>

          <div style="font-size:11px; color:#6b7280; margin-top:10px;">
            Thao tác có thể mất 10–60 giây tuỳ độ dài hội thoại và tốc độ model.
          </div>
        </div>
      </div>
    </div>
  `;

  const style = document.createElement("style");
  style.textContent = `@keyframes spin {{ to {{ transform: rotate(360deg); }} }}`;
  overlay.appendChild(style);

  let sec = 0;
  const timerEl = overlay.querySelector("#owui_modal_timer");
  const t = setInterval(() => {{
    sec += 1;
    if (timerEl) timerEl.textContent = String(sec);
  }}, 1000);
  overlay.dataset.timer = String(t);

  document.body.appendChild(overlay);
}})();
"""
        await __event_call__({"type": "execute", "data": {"code": js}})

    async def _ui_modal_update(
        self,
        __event_call__=None,
        title: Optional[str] = None,
        subtitle: Optional[str] = None,
        steps: Optional[List[str]] = None,
        done_ok: Optional[bool] = None,
        error_text: Optional[str] = None,
        fmt: str = "excel",
    ):
        if not __event_call__:
            return

        fmt_labels = {"pdf": "PDF", "excel": "Excel", "docx": "Word"}
        label = fmt_labels.get(fmt, "File")

        def _js_str(s: str) -> str:
            return s.replace("\\", "\\\\").replace("`", "\\`").replace("$", "\\$")

        steps_js = "null"
        if isinstance(steps, list):
            safe = [_js_str(x) for x in steps]
            steps_js = "[" + ",".join([f"`{x}`" for x in safe]) + "]"

        title_js = f"`{_js_str(title)}`" if isinstance(title, str) else "null"
        subtitle_js = f"`{_js_str(subtitle)}`" if isinstance(subtitle, str) else "null"
        error_js = f"`{_js_str(error_text)}`" if isinstance(error_text, str) else "null"
        done_js = (
            "true" if done_ok is True else ("false" if done_ok is False else "null")
        )

        js = f"""
(() => {{
  const overlay = document.getElementById("owui_export_modal");
  if (!overlay) return;

  const t = overlay.querySelector("#owui_modal_title");
  const st = overlay.querySelector("#owui_modal_subtitle");
  const stepsEl = overlay.querySelector("#owui_modal_steps");
  const closeBtn = overlay.querySelector("#owui_modal_close");
  const spinner = overlay.querySelector("#owui_modal_spinner");

  const title = {title_js};
  const subtitle = {subtitle_js};
  const steps = {steps_js};
  const done = {done_js};
  const err = {error_js};

  if (title) t.textContent = title;
  if (subtitle) st.textContent = subtitle;

  if (Array.isArray(steps) && stepsEl) {{
    stepsEl.innerHTML = steps.map(s => `<li>${{s}}</li>`).join("");
  }}

  if (done !== null) {{
    const tid = overlay.dataset.timer;
    if (tid) {{
      try {{ clearInterval(Number(tid)); }} catch (e) {{}}
    }}

    closeBtn.disabled = false;
    closeBtn.style.opacity = "1";
    closeBtn.style.cursor = "pointer";
    closeBtn.onclick = () => overlay.remove();

    if (spinner) {{
      spinner.style.animation = "none";
      spinner.style.borderTopColor = done ? "#16a34a" : "#dc2626";
      spinner.style.borderColor = done ? "#86efac" : "#fecaca";
      spinner.style.background = done ? "#dcfce7" : "#fee2e2";
    }}

    if (done === true) {{
      t.textContent = "Xuất {label} thành công";
      st.textContent = "File đang được tải xuống. Bạn có thể đóng cửa sổ này.";
    }} else {{
      t.textContent = "Xuất {label} thất bại";
      st.textContent = err ? err : "Có lỗi xảy ra. Vui lòng thử lại hoặc liên hệ quản trị.";
    }}
  }}
}})();
"""
        await __event_call__({"type": "execute", "data": {"code": js}})

    # =========================================================
    # Context helpers (from Excel tool - used as-is)
    # =========================================================

    def _get_current_model_id(self, body: dict, __model__=None) -> Optional[str]:
        m = body.get("model")
        if isinstance(m, str) and m.strip():
            return m.strip()
        if isinstance(__model__, dict):
            for k in ("id", "model", "name"):
                v = __model__.get(k)
                if isinstance(v, str) and v.strip():
                    return v.strip()
        return None

    def _get_messages_window(self, body: dict) -> List[Dict[str, Any]]:
        msgs = body.get("messages", [])
        if not isinstance(msgs, list):
            return []
        n = (
            int(self.valves.context_window_messages)
            if self.valves.context_window_messages
            else 0
        )
        if n <= 0:
            return msgs
        return msgs[-n:]

    def _sanitize_message_content(self, content: Any) -> str:
        if content is None:
            return ""

        if isinstance(content, str):
            s = content
        else:
            try:
                s = str(content)
            except Exception:
                s = ""

        if self.valves.drop_file_like_content:
            if "data:" in s and "base64" in s:
                s = re.sub(
                    r"data:[^\\s]+;base64,[A-Za-z0-9+/=]+", "[DATA_URL_REMOVED]", s
                )
            s = re.sub(r"[A-Za-z0-9+/=]{1500,}", "[BASE64_REMOVED]", s)

        if len(s) > 6000:
            s = s[:6000] + "…[TRUNCATED]"
        return s

    def _build_normalize_user_prompt(self, body: dict, fmt: str = "excel") -> str:
        mode = (self.valves.normalize_mode or "conversation_window").strip()

        if mode == "assistant_only":
            last_assistant = ""
            for m in reversed(body.get("messages", [])):
                if isinstance(m, dict) and m.get("role") == "assistant":
                    last_assistant = self._sanitize_message_content(m.get("content"))
                    if last_assistant.strip():
                        break
            if fmt == "excel":
                return (
                    "Hãy tạo bảng Markdown từ nội dung TRẢ LỜI CỦA BOT dưới đây.\n\n"
                    f"BOT:\n{last_assistant}"
                )
            else:
                return (
                    "Chuẩn hoá nội dung trả lời dưới đây thành Markdown đẹp.\n\n"
                    f"BOT:\n{last_assistant}"
                )

        # conversation_window (default)
        window = self._get_messages_window(body)
        lines = []
        for m in window:
            if not isinstance(m, dict):
                continue
            role = m.get("role", "")
            c = self._sanitize_message_content(m.get("content"))
            if not c.strip():
                continue
            if role == "system":
                lines.append(f"[SYSTEM]\n{c}")
            elif role == "user":
                lines.append(f"[USER]\n{c}")
            elif role == "assistant":
                lines.append(f"[ASSISTANT]\n{c}")
            else:
                lines.append(f"[{role.upper()}]\n{c}")

        if fmt == "excel":
            return (
                "Dưới đây là đoạn hội thoại gần nhất. Hãy dựa vào YÊU CẦU CỦA NGƯỜI DÙNG để tạo bảng Markdown từ nội dung phù hợp.\n"
                f"- Gợi ý: tối đa khoảng {int(self.valves.max_rows_hint)} dòng dữ liệu.\n\n"
                + "\n\n".join(lines)
            )
        elif fmt == "docx":
            return (
                "Dưới đây là hội thoại. Hãy dựa vào YÊU CẦU CỦA NGƯỜI DÙNG để tạo văn bản Markdown "
                "chuyên nghiệp, có cấu trúc rõ ràng (heading, list, table, bold, italic...).\n\n"
                + "\n\n".join(lines)
            )
        else:
            return (
                "Dưới đây là hội thoại. Tạo văn bản Markdown chuyên nghiệp từ nội dung.\n\n"
                + "\n\n".join(lines)
            )

    # =========================================================
    # Markdown table helpers (from Excel tool)
    # =========================================================

    def _is_table_line(self, line: str) -> bool:
        return line.count("|") >= 2 and len(line.strip()) >= 3

    def _is_separator_line(self, line: str) -> bool:
        s = line.strip().strip("|").strip()
        if not s:
            return False
        return all(ch in "-: |" for ch in line.strip()) and "-" in line

    def _split_row(self, line: str) -> List[str]:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    # =========================================================
    # OpenWebUI API (from Excel tool - used as-is)
    # =========================================================

    def _resolve_openwebui_base_url(self, __request__=None) -> Optional[str]:
        if (
            isinstance(self.valves.openwebui_base_url, str)
            and self.valves.openwebui_base_url.strip()
        ):
            return self.valves.openwebui_base_url.strip().rstrip("/")

        if __request__ is not None:
            try:
                base = str(__request__.base_url).rstrip("/")
                if base:
                    return base
            except Exception:
                pass

            try:
                origin = __request__.headers.get("origin")
                if origin:
                    return origin.rstrip("/")
                host = __request__.headers.get("host")
                proto = __request__.headers.get("x-forwarded-proto", "http")
                if host:
                    return f"{proto}://{host}".rstrip("/")
            except Exception:
                pass

        return None

    def _extract_auth_headers_from_request(self, __request__=None) -> Dict[str, str]:
        headers: Dict[str, str] = {"Content-Type": "application/json"}
        if __request__ is None:
            return headers

        try:
            auth = __request__.headers.get("authorization")
            if auth:
                headers["Authorization"] = auth
        except Exception:
            pass

        try:
            cookie = __request__.headers.get("cookie")
            if cookie and "Authorization" not in headers:
                headers["Cookie"] = cookie
        except Exception:
            pass

        return headers

    async def _requests_post_json(
        self, url: str, headers: Dict[str, str], payload: Dict[str, Any], timeout: int
    ):
        def _do():
            return requests.post(url, headers=headers, json=payload, timeout=timeout)

        return await asyncio.to_thread(_do)

    async def _normalize_markdown_via_openwebui(
        self,
        body: dict,
        __model__=None,
        __request__=None,
        fmt: str = "excel",
    ) -> str:
        base = self._resolve_openwebui_base_url(__request__)
        if not base:
            req_info = "no request object"
            if __request__:
                try:
                    req_info = f"host={__request__.headers.get('host')}, origin={__request__.headers.get('origin')}, base_url={__request__.base_url}"
                except Exception:
                    req_info = "headers unavailable"
            logger.error("Cannot resolve base_url. Request info: %s", req_info)
            raise ValueError(
                "Không xác định được Open WebUI base_url. Hãy set valves.openwebui_base_url (VD http://localhost:3000)."
            )

        url = base + self.valves.openwebui_chat_completions_path
        headers = self._extract_auth_headers_from_request(__request__)

        model_id = self._get_current_model_id(body, __model__)
        if not model_id:
            raise ValueError(
                "Không xác định được model đang dùng trong chat để gọi normalize."
            )

        user_prompt = self._build_normalize_user_prompt(body, fmt=fmt)

        # Pick instruction by format
        if fmt == "excel":
            sys_prompt = self.valves.normalize_instructions_excel
        elif fmt == "docx":
            sys_prompt = self.valves.normalize_instructions_docx
        else:
            sys_prompt = self.valves.normalize_instructions_pdf

        messages = [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ]

        payload = {"model": model_id, "messages": messages, "stream": False}
        resp = await self._requests_post_json(
            url, headers=headers, payload=payload, timeout=int(self.valves.timeout_sec)
        )

        if resp.status_code >= 400:
            raise ValueError(
                f"Normalize failed: HTTP {resp.status_code} - {resp.text[:500]}"
            )

        data = resp.json()
        content = None
        try:
            content = data.get("choices", [{}])[0].get("message", {}).get("content")
        except Exception:
            content = None

        if not content or not isinstance(content, str) or not content.strip():
            content = (
                data.get("message", {}).get("content") or data.get("content") or ""
            )

        if not content.strip():
            raise ValueError(
                "Normalize: không lấy được content từ /api/chat/completions response."
            )

        return content

    # =========================================================
    # Excel table parsing (from Excel tool)
    # =========================================================

    def _extract_tables(self, text: str) -> List[List[str]]:
        lines = text.splitlines()
        tables: List[List[str]] = []
        i = 0
        n = len(lines)
        while i < n:
            if (
                self._is_table_line(lines[i])
                and i + 1 < n
                and self._is_separator_line(lines[i + 1])
            ):
                block = [lines[i].strip(), lines[i + 1].strip()]
                i += 2
                while i < n and self._is_table_line(lines[i]):
                    block.append(lines[i].strip())
                    i += 1
                tables.append(block)
                continue
            i += 1
        return tables

    def _table_block_to_rows(self, block_lines: List[str]) -> List[List[str]]:
        if len(block_lines) < 2:
            return []
        header = self._split_row(block_lines[0])
        rows = [header]
        for ln in block_lines[2:]:
            rows.append(self._split_row(ln))

        max_cols = max(len(r) for r in rows) if rows else 0
        for r in rows:
            if len(r) < max_cols:
                r += [""] * (max_cols - len(r))
        return rows

    # =========================================================
    # Type inference (from Excel tool)
    # =========================================================

    _re_int = re.compile(r"^-?\d+$")
    _re_float = re.compile(r"^-?\d+(\.\d+)?$")
    _re_thousand = re.compile(r"^-?\d{1,3}([.,]\d{3})+(\.\d+)?$")
    _re_percent = re.compile(r"^-?\d+(\.\d+)?%$")
    _re_money = re.compile(r"^[₫$€]?\s*-?\d{1,3}([.,]\d{3})+(\.\d+)?\s*[₫$€]?$")
    _re_date1 = re.compile(r"^\d{4}-\d{2}-\d{2}$")
    _re_date2 = re.compile(r"^\d{2}/\d{2}/\d{4}$")

    def _parse_number(self, s: str) -> Optional[float]:
        t = s.strip().replace(" ", "")
        if self._re_percent.match(t):
            try:
                return float(t[:-1]) / 100.0
            except Exception:
                return None

        t2 = t.replace("₫", "").replace("$", "").replace("€", "")
        if self._re_thousand.match(t2) or self._re_money.match(t):
            if "," in t2 and "." in t2:
                t2 = t2.replace(",", "")
            else:
                t2 = t2.replace(".", "").replace(",", "")
            try:
                return float(t2)
            except Exception:
                return None

        if self._re_int.match(t2):
            try:
                return int(t2)
            except Exception:
                return None

        if self._re_float.match(t2):
            try:
                return float(t2)
            except Exception:
                return None

        return None

    def _maybe_parse_date(self, s: str):
        t = s.strip()
        try:
            import datetime as _dt

            if self._re_date1.match(t):
                y, m, d = t.split("-")
                return _dt.date(int(y), int(m), int(d))
            if self._re_date2.match(t):
                d, m, y = t.split("/")
                return _dt.date(int(y), int(m), int(d))
        except Exception:
            return None
        return None

    # =========================================================
    # Excel builder (from Excel tool - used as-is)
    # =========================================================

    def _autosize_columns(self, ws, min_w=10, max_w=55):
        max_w = int(self.valves.xlsx_max_col_width) if self.valves.xlsx_max_col_width else max_w
        for col_idx in range(1, ws.max_column + 1):
            max_len = 0
            for row in range(1, ws.max_row + 1):
                v = ws.cell(row=row, column=col_idx).value
                if v is None:
                    continue
                max_len = max(max_len, len(str(v)))
            width = min(max(min_w, max_len + 2), max_w)
            ws.column_dimensions[get_column_letter(col_idx)].width = width

    def _apply_borders(self, ws, start_row, start_col, end_row, end_col):
        thin = Side(style="thin", color="D9D9D9")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for r in range(start_row, end_row + 1):
            for c in range(start_col, end_col + 1):
                ws.cell(r, c).border = border

    def _build_xlsx_bytes(self, markdown_text: str) -> bytes:
        wb = Workbook()
        ws = wb.active
        ws.title = self.valves.xlsx_sheet_name

        ws["A1"] = self.valves.xlsx_title
        ws["A1"].font = Font(bold=True, size=16)
        ws["A2"] = (
            f"Thời gian tạo: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        ws["A2"].font = Font(size=10, color="666666")

        tables = self._extract_tables(markdown_text)
        chosen = (
            tables[self.valves.xlsx_table_index]
            if len(tables) > self.valves.xlsx_table_index
            else None
        )
        start_row = int(self.valves.xlsx_table_start_row)

        if not chosen:
            ws[f"A{start_row}"] = (
                "Không tìm thấy Markdown table trong nội dung để xuất."
            )
            ws[f"A{start_row}"].font = Font(bold=True, color="C00000")
            ws[f"A{start_row+2}"] = "raw_text"
            ws[f"A{start_row+2}"].font = Font(bold=True)
            ws[f"A{start_row+3}"] = markdown_text
            ws[f"A{start_row+3}"].alignment = Alignment(wrap_text=True, vertical="top")
            ws.column_dimensions["A"].width = 90

            bio = io.BytesIO()
            wb.save(bio)
            bio.seek(0)
            return bio.getvalue()

        rows = self._table_block_to_rows(chosen)
        if not rows or len(rows) < 1:
            raise ValueError("Bảng Markdown không hợp lệ (không có header).")

        headers = rows[0]
        data_rows = rows[1:]

        for c_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=start_row, column=c_idx, value=h)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E79")
            cell.alignment = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )

        for r_off, row in enumerate(data_rows, start=1):
            for c_idx, v in enumerate(row, start=1):
                val = v.strip() if isinstance(v, str) else (v if v is not None else "")

                d = self._maybe_parse_date(val) if isinstance(val, str) else None
                if d is not None:
                    cell = ws.cell(row=start_row + r_off, column=c_idx, value=d)
                    cell.number_format = "yyyy-mm-dd"
                else:
                    num = self._parse_number(val) if isinstance(val, str) else None
                    if num is not None:
                        cell = ws.cell(row=start_row + r_off, column=c_idx, value=num)
                        if isinstance(val, str) and val.strip().endswith("%"):
                            cell.number_format = "0.00%"
                        else:
                            cell.number_format = (
                                "#,##0.00"
                                if (
                                    isinstance(num, float)
                                    and not float(num).is_integer()
                                )
                                else "#,##0"
                            )
                    else:
                        cell = ws.cell(row=start_row + r_off, column=c_idx, value=val)

                cell.alignment = Alignment(vertical="top", wrap_text=True)

        end_row = start_row + len(data_rows)
        end_col = len(headers)

        ws.freeze_panes = ws.cell(row=start_row + 1, column=1)
        self._apply_borders(ws, start_row, 1, end_row, end_col)

        table_ref = f"A{start_row}:{get_column_letter(end_col)}{end_row}"
        tname = f"T_{datetime.datetime.now().strftime('%H%M%S')}"
        tab = Table(displayName=tname, ref=table_ref)
        tab.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        ws.add_table(tab)

        self._autosize_columns(ws)
        ws.row_dimensions[start_row].height = 22

        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)
        return bio.getvalue()

    # =========================================================
    # PDF builder (from PDF tool - used as-is)
    # =========================================================

    def _render_inline_text_pdf(self, pdf, text):
        """Parse and render inline Markdown (bold, italic, code) using write()."""
        parts = []
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'
            r'|(\*\*(.+?)\*\*)'
            r'|(\*(.+?)\*)'
            r'|(`([^`]+)`)'
            r'|([^*`]+)'
        )
        for m in pattern.finditer(text):
            if m.group(2):
                parts.append(("BI", m.group(2)))
            elif m.group(4):
                parts.append(("B", m.group(4)))
            elif m.group(6):
                parts.append(("I", m.group(6)))
            elif m.group(8):
                parts.append(("CODE", m.group(8)))
            elif m.group(9):
                parts.append(("", m.group(9)))

        for style, txt in parts:
            if style == "CODE":
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body - 1)
                pdf.set_text_color(199, 37, 78)
                pdf.write(5, txt)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
            else:
                pdf.set_font(pdf._font_name, style, self.valves.pdf_font_size_body)
                pdf.write(5, txt)
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)

    def _build_pdf_bytes(self, markdown_text: str) -> bytes:
        pdf = MarkdownPDF(
            font_size=self.valves.pdf_font_size_body,
            margin=self.valves.pdf_page_margin,
        )
        pdf.alias_nb_pages()
        pdf.add_page()
        pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)

        lines = markdown_text.splitlines()
        i = 0
        n = len(lines)
        in_code_block = False
        code_lines = []

        while i < n:
            line = lines[i]
            stripped = line.strip()

            # Code block
            if stripped.startswith("```"):
                if in_code_block:
                    pdf.set_fill_color(245, 245, 245)
                    pdf.set_font(pdf._font_name, "", 9)
                    pdf.set_text_color(30, 30, 30)
                    code_text = "\n".join(code_lines)
                    w = pdf.w - pdf.l_margin - pdf.r_margin
                    pdf.multi_cell(w, 4.5, code_text, border=1, fill=True)
                    pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(3)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Empty line
            if not stripped:
                pdf.ln(3)
                i += 1
                continue

            # Headings
            if stripped.startswith("### "):
                pdf.ln(4)
                pdf.set_font(pdf._font_name, "B", self.valves.pdf_font_size_h3)
                pdf.set_text_color(31, 78, 121)
                pdf.multi_cell(0, 7, stripped[4:])
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(2)
                i += 1
                continue

            if stripped.startswith("## "):
                pdf.ln(5)
                pdf.set_font(pdf._font_name, "B", self.valves.pdf_font_size_h2)
                pdf.set_text_color(31, 78, 121)
                pdf.multi_cell(0, 7, stripped[3:])
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(2)
                i += 1
                continue

            if stripped.startswith("# "):
                pdf.ln(6)
                pdf.set_font(pdf._font_name, "B", self.valves.pdf_font_size_h1)
                pdf.set_text_color(13, 55, 107)
                pdf.multi_cell(0, 9, stripped[2:])
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(3)
                i += 1
                continue

            # Blockquote
            if stripped.startswith("> "):
                pdf.set_text_color(85, 85, 85)
                pdf.set_font(pdf._font_name, "I", self.valves.pdf_font_size_body)
                pdf.set_x(pdf.l_margin + 10)
                pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 10, 5, stripped[2:])
                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(2)
                i += 1
                continue

            # Horizontal rule
            if stripped in ("---", "***", "___"):
                y = pdf.get_y()
                pdf.set_draw_color(200, 200, 200)
                pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                pdf.ln(5)
                i += 1
                continue

            # Table
            if (self._is_table_line(stripped)
                    and i + 1 < n
                    and self._is_separator_line(lines[i + 1].strip())):
                headers = self._split_row(stripped)
                i += 2
                data_rows = []
                while i < n and self._is_table_line(lines[i].strip()):
                    data_rows.append(self._split_row(lines[i].strip()))
                    i += 1

                num_cols = len(headers)
                if num_cols == 0:
                    continue

                avail_w = pdf.w - pdf.l_margin - pdf.r_margin
                col_w = avail_w / num_cols

                pdf.set_font(pdf._font_name, "B", self.valves.pdf_font_size_body)
                pdf.set_fill_color(31, 78, 121)
                pdf.set_text_color(255, 255, 255)
                for ci, h in enumerate(headers):
                    pdf.cell(col_w, 7, h[:30], border=1, fill=True, align="C")
                pdf.ln()

                pdf.set_font(pdf._font_name, "", self.valves.pdf_font_size_body)
                pdf.set_text_color(0, 0, 0)
                for ri, row in enumerate(data_rows):
                    if ri % 2 == 0:
                        pdf.set_fill_color(242, 247, 251)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    for ci in range(num_cols):
                        val = row[ci] if ci < len(row) else ""
                        pdf.cell(col_w, 6, val[:40], border=1, fill=True)
                    pdf.ln()

                pdf.ln(3)
                continue

            # Unordered list
            if re.match(r'^[\-\*\+]\s', stripped):
                text = re.sub(r'^[\-\*\+]\s+', '', stripped)
                pdf.set_x(pdf.l_margin + 5)
                pdf.write(5, "• ")
                self._render_inline_text_pdf(pdf, text)
                pdf.ln(6)
                i += 1
                continue

            # Ordered list
            ol_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if ol_match:
                num = ol_match.group(1)
                text = ol_match.group(2)
                pdf.set_x(pdf.l_margin + 5)
                pdf.write(5, f"{num}. ")
                self._render_inline_text_pdf(pdf, text)
                pdf.ln(6)
                i += 1
                continue

            # Regular paragraph
            self._render_inline_text_pdf(pdf, stripped)
            pdf.ln(6)
            i += 1

        # Footer timestamp
        pdf.ln(10)
        pdf.set_font(pdf._font_name, "I", 8)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 5,
                 f"Tạo lúc: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                 align="R")

        raw = pdf.output()
        return bytes(raw) if not isinstance(raw, bytes) else raw

    # =========================================================
    # DOCX builder (from DOCX tool - used as-is)
    # =========================================================

    def _set_cell_shading(self, cell, color_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color_hex)
        tcPr.append(shading)

    def _add_formatted_run(self, paragraph, text, bold=False, italic=False, code=False):
        run = paragraph.add_run(text)
        run.font.name = self.valves.docx_font_name if not code else "Consolas"
        run.font.size = Pt(self.valves.docx_font_size_body)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.size = Pt(self.valves.docx_font_size_body - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        return run

    def _parse_inline(self, paragraph, text):
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'
            r'|(\*\*(.+?)\*\*)'
            r'|(\*(.+?)\*)'
            r'|(`([^`]+)`)'
            r'|([^*`]+)'
        )
        for m in pattern.finditer(text):
            if m.group(2):
                self._add_formatted_run(paragraph, m.group(2), bold=True, italic=True)
            elif m.group(4):
                self._add_formatted_run(paragraph, m.group(4), bold=True)
            elif m.group(6):
                self._add_formatted_run(paragraph, m.group(6), italic=True)
            elif m.group(8):
                self._add_formatted_run(paragraph, m.group(8), code=True)
            elif m.group(9):
                self._add_formatted_run(paragraph, m.group(9))

    def _build_docx_bytes(self, markdown_text: str) -> bytes:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(self.valves.docx_page_margin_cm)
            section.bottom_margin = Cm(self.valves.docx_page_margin_cm)
            section.left_margin = Cm(self.valves.docx_page_margin_cm)
            section.right_margin = Cm(self.valves.docx_page_margin_cm)

        style = doc.styles["Normal"]
        style.font.name = self.valves.docx_font_name
        style.font.size = Pt(self.valves.docx_font_size_body)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        lines = markdown_text.splitlines()
        i = 0
        n = len(lines)
        in_code_block = False
        code_lines = []

        while i < n:
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    for ci, cl in enumerate(code_lines):
                        run = p.add_run(cl)
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                        if ci < len(code_lines) - 1:
                            p.add_run("\n").font.size = Pt(10)
                    pPr = p._p.get_or_add_pPr()
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:val"), "clear")
                    shading.set(qn("w:fill"), "F5F5F5")
                    pPr.append(shading)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            if not stripped:
                i += 1
                continue

            if stripped.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[4:])
                run.bold = True
                run.font.size = Pt(self.valves.docx_font_size_h3)
                run.font.name = self.valves.docx_font_name
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                p.paragraph_format.space_before = Pt(12)
                i += 1
                continue

            if stripped.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[3:])
                run.bold = True
                run.font.size = Pt(self.valves.docx_font_size_h2)
                run.font.name = self.valves.docx_font_name
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                p.paragraph_format.space_before = Pt(14)
                i += 1
                continue

            if stripped.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:])
                run.bold = True
                run.font.size = Pt(self.valves.docx_font_size_h1)
                run.font.name = self.valves.docx_font_name
                run.font.color.rgb = RGBColor(0x0D, 0x37, 0x6B)
                p.paragraph_format.space_before = Pt(18)
                i += 1
                continue

            if stripped.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                self._parse_inline(p, stripped[2:])
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                i += 1
                continue

            if stripped in ("---", "***", "___"):
                p = doc.add_paragraph()
                p.add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                i += 1
                continue

            if (self._is_table_line(stripped)
                    and i + 1 < n
                    and self._is_separator_line(lines[i + 1].strip())):
                headers = self._split_row(stripped)
                i += 2
                data_rows = []
                while i < n and self._is_table_line(lines[i].strip()):
                    data_rows.append(self._split_row(lines[i].strip()))
                    i += 1

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"

                for ci, h in enumerate(headers):
                    cell = table.rows[0].cells[ci] if ci < num_cols else None
                    if cell:
                        cell.text = h
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.bold = True
                                run.font.size = Pt(self.valves.docx_font_size_body)
                                run.font.name = self.valves.docx_font_name
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        self._set_cell_shading(cell, "1F4E79")

                for ri, row in enumerate(data_rows):
                    for ci in range(num_cols):
                        cell = table.rows[ri + 1].cells[ci]
                        val = row[ci] if ci < len(row) else ""
                        cell.text = val
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(self.valves.docx_font_size_body)
                                run.font.name = self.valves.docx_font_name
                        if ri % 2 == 0:
                            self._set_cell_shading(cell, "F2F7FB")
                continue

            if re.match(r'^[\-\*\+]\s', stripped):
                text = re.sub(r'^[\-\*\+]\s+', '', stripped)
                p = doc.add_paragraph(style="List Bullet")
                self._parse_inline(p, text)
                i += 1
                continue

            ol_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if ol_match:
                text = ol_match.group(2)
                p = doc.add_paragraph(style="List Number")
                self._parse_inline(p, text)
                i += 1
                continue

            p = doc.add_paragraph()
            self._parse_inline(p, stripped)
            i += 1

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Tạo lúc: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = self.valves.docx_font_name

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    # =========================================================
    # Download trigger (JS injection)
    # =========================================================

    async def _trigger_download(
        self, __event_call__, file_bytes: bytes, filename: str, mime: str
    ):
        if not __event_call__:
            return
        if len(file_bytes) > 4_000_000:
            logger.warning(
                "Large file (%d bytes / ~%.1f MB). JS download may be slow.",
                len(file_bytes), len(file_bytes) / 1_048_576
            )
        b64 = base64.b64encode(file_bytes).decode("ascii")
        js = f"""
try {{
  const b64 = "{b64}";
  const raw = atob(b64);
  const arr = new Uint8Array(raw.length);
  for (let i = 0; i < raw.length; i++) arr[i] = raw.charCodeAt(i);
  const blob = new Blob([arr], {{ type: "{mime}" }});
  const url = URL.createObjectURL(blob);
  const a = document.createElement("a");
  a.href = url;
  a.download = "{filename}";
  document.body.appendChild(a);
  a.click();
  setTimeout(() => {{ URL.revokeObjectURL(url); a.remove(); }}, 3000);
}} catch (e) {{
  console.error("Download failed", e);
}}
"""
        await __event_call__({"type": "execute", "data": {"code": js}})

    # =========================================================
    # Fallback: raw Markdown from last assistant message
    # =========================================================

    def _fallback_raw_markdown(self, body: dict) -> str:
        for m in reversed(body.get("messages", [])):
            if isinstance(m, dict) and m.get("role") == "assistant":
                c = m.get("content", "")
                if isinstance(c, str) and c.strip():
                    return c.strip()
        return ""

    # =========================================================
    # MAIN ACTION
    # =========================================================

    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__=None,
        __model__=None,
        __request__=None,
        __id__=None,
    ) -> Optional[dict]:
        logger.info("=== Export All: action() called ===")

        # 1. Show format picker
        fmt = await self._show_format_picker(__event_call__)
        if not fmt:
            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": "Đã huỷ xuất file.",
                            "done": True,
                        },
                    }
                )
            return None

        fmt_labels = {"pdf": "PDF", "excel": "Excel", "docx": "Word"}
        label = fmt_labels.get(fmt, fmt)
        logger.info(f"User selected format: {fmt}")

        # 2. Open progress modal
        await self._ui_modal_open(__event_call__, fmt=fmt)

        try:
            # 3. Normalize via LLM
            await self._ui_modal_update(
                __event_call__,
                subtitle=f"Đang chuẩn hoá nội dung cho {label}...",
                steps=[
                    "✅ Chọn định dạng: " + label,
                    "⏳ Chuẩn hoá nội dung (LLM)...",
                ],
                fmt=fmt,
            )

            markdown_text = ""
            if self.valves.normalize_with_llm:
                try:
                    markdown_text = await self._normalize_markdown_via_openwebui(
                        body, __model__=__model__, __request__=__request__, fmt=fmt
                    )
                except Exception as e:
                    logger.warning(f"LLM normalize failed, using fallback: {e}")
                    markdown_text = self._fallback_raw_markdown(body)
            else:
                markdown_text = self._fallback_raw_markdown(body)

            if not markdown_text.strip():
                raise ValueError("Không có nội dung để xuất.")

            # 4. Build file
            await self._ui_modal_update(
                __event_call__,
                subtitle=f"Đang tạo file {label}...",
                steps=[
                    "✅ Chọn định dạng: " + label,
                    "✅ Chuẩn hoá nội dung",
                    f"⏳ Tạo file {label}...",
                ],
                fmt=fmt,
            )

            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

            if fmt == "pdf":
                file_bytes = self._build_pdf_bytes(markdown_text)
                filename = f"{ts}_{self.valves.pdf_base_filename}"
                mime = "application/pdf"
            elif fmt == "excel":
                file_bytes = self._build_xlsx_bytes(markdown_text)
                filename = f"{ts}_{self.valves.xlsx_base_filename}"
                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            elif fmt == "docx":
                file_bytes = self._build_docx_bytes(markdown_text)
                filename = f"{ts}_{self.valves.docx_base_filename}"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                raise ValueError(f"Định dạng không hỗ trợ: {fmt}")

            # 5. Download
            await self._ui_modal_update(
                __event_call__,
                subtitle="Đang tải file xuống...",
                steps=[
                    "✅ Chọn định dạng: " + label,
                    "✅ Chuẩn hoá nội dung",
                    f"✅ Tạo file {label}",
                    "⏳ Tải file xuống...",
                ],
                fmt=fmt,
            )

            await self._trigger_download(__event_call__, file_bytes, filename, mime)

            # 6. Done
            await asyncio.sleep(self.valves.min_wizard_seconds)
            await self._ui_modal_update(
                __event_call__,
                steps=[
                    "✅ Chọn định dạng: " + label,
                    "✅ Chuẩn hoá nội dung",
                    f"✅ Tạo file {label}",
                    "✅ Tải file xuống",
                ],
                done_ok=True,
                fmt=fmt,
            )

            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": f"✅ Đã xuất {label} thành công!",
                            "done": True,
                        },
                    }
                )

            logger.info(f"Export {fmt} completed successfully.")
            return body

        except Exception as exc:
            logger.exception(f"Export {fmt} failed")
            await self._ui_modal_update(
                __event_call__,
                error_text=str(exc)[:300],
                done_ok=False,
                fmt=fmt,
            )
            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": f"❌ Xuất {label} thất bại: {str(exc)[:100]}",
                            "done": True,
                        },
                    }
                )

            return None
